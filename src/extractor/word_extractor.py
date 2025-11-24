from docx import Document

from .extractor_base import ExtractorBase


class WordExtractor(ExtractorBase):
    def extract(self) -> str:
        doc = Document(str(self.path))
        lines = []

        # 抽出: 段落
        for para in doc.paragraphs:
            lines.append(para.text)

        # 抽出: テーブル（段落の後に追加）
        if doc.tables:
            lines.append("")
            for ti, table in enumerate(doc.tables, start=1):
                lines.append(f"[Table {ti}]")
                for row in table.rows:
                    # 各セルのテキストをタブ区切りで出力
                    lines.append("\t".join([cell.text for cell in row.cells]))
                lines.append("")

        return "\n".join(lines)
